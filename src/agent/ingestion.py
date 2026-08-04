"""Document ingestion pipeline.

Parses documents (docx, xlsx), extracts metadata from filenames,
extracts body text, and builds a ChromaDB index for semantic search.
"""

from __future__ import annotations

import json
import logging
import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Protocol

import chromadb
from docx import Document as DocxDocument
from openpyxl import load_workbook

from agent.utils import (
    CHROMA_MAX_CHARS,
    CHUNK_MAX_CHARS,
    DOC_TYPE_LABELS,
    INDEX_BATCH_SIZE,
    MIN_SECTION_CHARS,
    BodyStoreEntry,
    make_unique_id,
)

logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    """Represents a parsed document with metadata and body text."""

    doc_id: str
    doc_type: str
    doc_type_label: str
    title: str
    revision: str
    is_obsolete: bool
    is_signed: bool
    file_format: str
    filename: str
    filepath: str
    body: str
    sections: list[str] = field(default_factory=list)


class MetadataExtractor(Protocol):
    """Interface for extracting document metadata from filenames.

    Implement this protocol to support different filename conventions.
    The default implementation is QMSFilenameExtractor, which uses a regex.
    Future implementations could use lightweight LLM-based extraction for more complex patterns.
    """

    def parse(self, filename: str) -> dict | None:
        """Extract metadata from a filename.

        Returns a dict with keys: doc_id, doc_type, doc_type_label, title,
        revision, is_obsolete, is_signed, file_format, filename.
        Returns None if the filename doesn't match the expected pattern.
        """
        ...


class QMSFilenameExtractor:
    """Extracts metadata from QMS filenames following the pattern:
    {TYPE}-{ID} - {TITLE}_{REVISION}[-STATUS][.ext]

    Handles edge cases: 3P-P01-32 (short number), IFU-MX1 (alpha suffix),
    VVPR-P01-152- (dash glued), MEMO-P01-660 (space), and status suffix spacing.
    """

    FILENAME_RE = re.compile(
        r"^(?P<doc_id>[A-Z0-9]{2,4}-(?:[A-Z0-9]+-)*[A-Z0-9]+)"
        r"[\s-]+"
        r"(?P<title>.+?)"
        r"_(?P<revision>[A-Z])"
        r"(?P<status>[\s_-]*(?:signed|Signed|Obsolete|obsolete))?"
        r"(?:\(\d+\))?"
        r"(?:\.(?P<ext>docx|xlsx))?$"
    )

    def __init__(self, type_labels: dict[str, str] | None = None):
        self._type_labels = type_labels or DOC_TYPE_LABELS

    def parse(self, filename: str) -> dict | None:
        match = self.FILENAME_RE.match(filename)
        if not match:
            logger.warning("Could not parse filename: %s", filename)
            return None

        doc_id = match.group("doc_id")
        title = match.group("title").strip()
        revision = match.group("revision")
        status = match.group("status") or ""
        ext = match.group("ext") or "unknown"

        doc_type = doc_id.split("-")[0]
        doc_type_label = self._type_labels.get(doc_type)
        if doc_type_label is None:
            logger.warning(
                "Unknown document type prefix: %s (from %s)", doc_type, filename
            )
            doc_type_label = f"Unknown ({doc_type})"

        return {
            "doc_id": doc_id,
            "doc_type": doc_type,
            "doc_type_label": doc_type_label,
            "title": title,
            "revision": revision,
            "is_obsolete": "obsolete" in status.lower(),
            "is_signed": "signed" in status.lower(),
            "file_format": ext,
            "filename": filename,
        }


def parse_filename(filename: str) -> dict | None:
    """Parse a QMS filename using the default extractor. Convenience wrapper."""
    return QMSFilenameExtractor().parse(filename)


def extract_docx_text(filepath: str) -> tuple[str, list[str]]:
    """Extract text from a .docx file. Returns (full_body, sections).

    Sections are split on headings. Tables are appended as separate sections.
    """
    try:
        doc = DocxDocument(filepath)
    except Exception as e:
        logger.error("Failed to parse docx %s: %s", filepath, e)
        return "", []

    paragraphs: list[str] = []
    sections: list[str] = []
    current_section_lines: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        paragraphs.append(text)

        style_name = para.style.name if para.style else ""
        is_heading = style_name.startswith("Heading") or style_name.startswith("Title")

        if is_heading:
            if current_section_lines:
                sections.append("\n".join(current_section_lines))
            current_section_lines = [text]
        else:
            current_section_lines.append(text)

    if current_section_lines:
        sections.append("\n".join(current_section_lines))

    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            row_text = " | ".join(cells)
            if row_text.strip(" |"):
                table_rows.append(row_text)
        if table_rows:
            table_text = "\n".join(table_rows)
            paragraphs.append(table_text)
            sections.append(table_text)

    full_body = "\n\n".join(paragraphs)

    if not sections and full_body:
        sections = _chunk_text(full_body)

    return full_body, sections


def extract_xlsx_text(filepath: str) -> tuple[str, list[str]]:
    """Extract text from an .xlsx file. Returns (full_body, sections).

    Each sheet becomes its own section.
    """
    try:
        wb = load_workbook(filepath, read_only=True, data_only=True)
    except Exception as e:
        logger.error("Failed to parse xlsx %s: %s", filepath, e)
        return "", []

    all_text_parts: list[str] = []
    sections: list[str] = []

    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        sheet_lines = [f"[Sheet: {sheet_name}]"]

        for row in sheet.iter_rows(values_only=True):
            cells = [str(c).strip() if c is not None else "" for c in row]
            row_text = " | ".join(c for c in cells if c)
            if row_text:
                sheet_lines.append(row_text)

        sheet_text = "\n".join(sheet_lines)
        all_text_parts.append(sheet_text)
        sections.append(sheet_text)

    wb.close()
    full_body = "\n\n".join(all_text_parts)

    if not sections and full_body:
        sections = _chunk_text(full_body)

    return full_body, sections


def _chunk_text(text: str, max_chars: int = CHUNK_MAX_CHARS) -> list[str]:
    """Split text into roughly equal chunks at paragraph boundaries."""
    paragraphs = text.split("\n\n")
    chunks: list[str] = []
    current_chunk: list[str] = []
    current_len = 0

    for para in paragraphs:
        if current_len + len(para) > max_chars and current_chunk:
            chunks.append("\n\n".join(current_chunk))
            current_chunk = []
            current_len = 0
        current_chunk.append(para)
        current_len += len(para)

    if current_chunk:
        chunks.append("\n\n".join(current_chunk))

    return chunks


def parse_all_documents(
    data_dir: str,
    extractor: MetadataExtractor | None = None,
) -> list[ParsedDocument]:
    """Parse all documents in the given directory.

    Args:
        data_dir: Path to directory containing document files.
        extractor: Metadata extractor to use. Defaults to QMSFilenameExtractor.
    """
    data_path = Path(data_dir)
    if not data_path.is_dir():
        raise FileNotFoundError(f"Data directory not found: {data_dir}")

    if extractor is None:
        extractor = QMSFilenameExtractor()

    docs: list[ParsedDocument] = []

    for entry in sorted(data_path.iterdir()):
        meta = extractor.parse(entry.name)
        if meta is None:
            continue

        filepath = str(entry)
        body = ""
        sections: list[str] = []

        if meta["file_format"] == "docx":
            body, sections = extract_docx_text(filepath)
        elif meta["file_format"] == "xlsx":
            body, sections = extract_xlsx_text(filepath)

        docs.append(
            ParsedDocument(
                doc_id=meta["doc_id"],
                doc_type=meta["doc_type"],
                doc_type_label=meta["doc_type_label"],
                title=meta["title"],
                revision=meta["revision"],
                is_obsolete=meta["is_obsolete"],
                is_signed=meta["is_signed"],
                file_format=meta["file_format"],
                filename=meta["filename"],
                filepath=filepath,
                body=body,
                sections=sections,
            )
        )

    logger.info("Parsed %d documents from %s", len(docs), data_dir)
    return docs


def determine_latest_revisions(docs: list[ParsedDocument]) -> dict[str, str]:
    """For each doc_id, determine which revision letter is the latest (highest non-obsolete)."""
    id_revisions: dict[str, list[tuple[str, bool]]] = {}
    for doc in docs:
        id_revisions.setdefault(doc.doc_id, []).append((doc.revision, doc.is_obsolete))

    latest: dict[str, str] = {}
    for doc_id, revisions in id_revisions.items():
        non_obsolete = [(r, obs) for r, obs in revisions if not obs]
        if non_obsolete:
            latest[doc_id] = max(non_obsolete, key=lambda x: x[0])[0]
        else:
            latest[doc_id] = max(revisions, key=lambda x: x[0])[0]

    return latest


def _dedup_id(base_id: str, seen: set[str]) -> str:
    """Append a suffix if base_id already exists in seen."""
    if base_id not in seen:
        seen.add(base_id)
        return base_id
    n = 2
    while f"{base_id}_{n}" in seen:
        n += 1
    deduped = f"{base_id}_{n}"
    seen.add(deduped)
    return deduped


def build_index(
    docs: list[ParsedDocument],
    persist_dir: str = ".chroma_index",
) -> tuple[chromadb.ClientAPI, dict[str, str]]:
    """Build a ChromaDB sections collection and return (client, latest_map).

    Only creates one collection ("sections") for semantic/keyword search.
    Document-level metadata lives in the body store JSON.
    """
    latest_map = determine_latest_revisions(docs)
    client = chromadb.PersistentClient(path=persist_dir)

    try:
        client.delete_collection("sections")
    except Exception:
        pass

    section_collection = client.create_collection(
        name="sections",
        metadata={"hnsw:space": "cosine"},
    )

    section_ids: list[str] = []
    section_documents: list[str] = []
    section_metadatas: list[dict] = []
    seen_ids: set[str] = set[str]()

    for doc in docs:
        is_latest = latest_map.get(doc.doc_id) == doc.revision
        unique_id = _dedup_id(make_unique_id(doc.doc_id, doc.revision), seen_ids)

        for i, section in enumerate(doc.sections):
            # Skip tiny sections -- assumed to be noise.
            if len(section) < MIN_SECTION_CHARS:
                continue

            section_embed = (
                f"[{doc.doc_id} Rev {doc.revision}: {doc.title} | {doc.doc_type_label}]\n"
                f"{section}"
            )

            section_ids.append(f"{unique_id}_sec{i}")
            section_documents.append(section_embed[:CHROMA_MAX_CHARS])
            section_metadatas.append(
                {
                    "doc_id": doc.doc_id,
                    "doc_type": doc.doc_type,
                    "title": doc.title,
                    "revision": doc.revision,
                    "is_latest": is_latest,
                    "is_obsolete": doc.is_obsolete,
                    "parent_unique_id": unique_id,
                    "section_index": i,
                }
            )

    for i in range(0, len(section_ids), INDEX_BATCH_SIZE):
        end = min(i + INDEX_BATCH_SIZE, len(section_ids))
        section_collection.add(
            ids=section_ids[i:end],
            documents=section_documents[i:end],
            metadatas=section_metadatas[i:end],
        )
        logger.info("Indexed sections batch %d-%d / %d", i + 1, end, len(section_ids))

    logger.info(
        "Indexed %d documents (%d sections) into ChromaDB at %s",
        len(docs),
        len(section_ids),
        persist_dir,
    )
    return client, latest_map


def _body_store_path(persist_dir: str) -> str:
    return os.path.join(persist_dir, "body_store.json")


def save_body_store(
    docs: list[ParsedDocument],
    latest_map: dict[str, str],
    persist_dir: str = ".chroma_index",
) -> None:
    """Save full document bodies + metadata to a JSON file."""
    os.makedirs(persist_dir, exist_ok=True)
    store: dict[str, BodyStoreEntry] = {}
    for doc in docs:
        uid = make_unique_id(doc.doc_id, doc.revision)
        store[uid] = {
            "doc_id": doc.doc_id,
            "doc_type": doc.doc_type,
            "doc_type_label": doc.doc_type_label,
            "title": doc.title,
            "revision": doc.revision,
            "is_latest": latest_map.get(doc.doc_id) == doc.revision,
            "is_obsolete": doc.is_obsolete,
            "is_signed": doc.is_signed,
            "file_format": doc.file_format,
            "filename": doc.filename,
            "has_body": bool(doc.body),
            "body": doc.body,
        }
    path = _body_store_path(persist_dir)
    with open(path, "w") as f:
        json.dump(store, f)
    logger.info("Saved body store (%d docs) to %s", len(store), path)


def ingest(
    data_dir: str,
    persist_dir: str = ".chroma_index",
    extractor: MetadataExtractor | None = None,
) -> chromadb.ClientAPI:
    """Full ingestion pipeline: parse, index, and save body store.

    Args:
        data_dir: Path to document directory.
        persist_dir: Path to store the ChromaDB index and body store.
        extractor: Metadata extractor for filename parsing. Defaults to QMS format.
    """
    logger.info("Starting ingestion from %s ...", data_dir)
    docs = parse_all_documents(data_dir, extractor=extractor)
    client, latest_map = build_index(docs, persist_dir=persist_dir)
    save_body_store(docs, latest_map=latest_map, persist_dir=persist_dir)
    logger.info("Ingestion complete.")
    return client
