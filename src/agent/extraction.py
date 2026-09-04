"""Document text extraction.

Kept free of heavy imports (ChromaDB, LangChain) so that parsing worker
processes start quickly: on spawn platforms such as Windows, every worker
re-imports the module that defines its task function.
"""

from __future__ import annotations

import logging

from docx import Document as DocxDocument
from openpyxl import load_workbook

from agent.utils import CHUNK_MAX_CHARS

logger = logging.getLogger(__name__)


def extract_docx_text(filepath: str) -> tuple[str, list[str]]:
    """Extract text from a .docx file. Returns (full_body, sections).

    Sections are split on headings. Tables are appended as separate sections.
    """
    try:
        doc = DocxDocument(filepath)
    except Exception as e:
        logger.error("Failed to parse docx %s: %s", filepath, e)
        return "", []

    paragraphs: list[str] = []
    sections: list[str] = []
    current_section_lines: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        paragraphs.append(text)

        style_name = para.style.name if para.style else ""
        is_heading = style_name.startswith("Heading") or style_name.startswith("Title")

        if is_heading:
            if current_section_lines:
                sections.append("\n".join(current_section_lines))
            current_section_lines = [text]
        else:
            current_section_lines.append(text)

    if current_section_lines:
        sections.append("\n".join(current_section_lines))

    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            row_text = " | ".join(cells)
            if row_text.strip(" |"):
                table_rows.append(row_text)
        if table_rows:
            table_text = "\n".join(table_rows)
            paragraphs.append(table_text)
            sections.append(table_text)

    full_body = "\n\n".join(paragraphs)

    if not sections and full_body:
        sections = chunk_text(full_body)

    return full_body, sections


def extract_xlsx_text(filepath: str) -> tuple[str, list[str]]:
    """Extract text from an .xlsx file. Returns (full_body, sections).

    Each sheet becomes its own section.
    """
    try:
        wb = load_workbook(filepath, read_only=True, data_only=True)
    except Exception as e:
        logger.error("Failed to parse xlsx %s: %s", filepath, e)
        return "", []

    all_text_parts: list[str] = []
    sections: list[str] = []

    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        sheet_lines = [f"[Sheet: {sheet_name}]"]

        for row in sheet.iter_rows(values_only=True):
            cells = [str(c).strip() if c is not None else "" for c in row]
            row_text = " | ".join(c for c in cells if c)
            if row_text:
                sheet_lines.append(row_text)

        sheet_text = "\n".join(sheet_lines)
        all_text_parts.append(sheet_text)
        sections.append(sheet_text)

    wb.close()
    full_body = "\n\n".join(all_text_parts)

    if not sections and full_body:
        sections = chunk_text(full_body)

    return full_body, sections


def chunk_text(text: str, max_chars: int = CHUNK_MAX_CHARS) -> list[str]:
    """Split text into roughly equal chunks at paragraph boundaries."""
    paragraphs = text.split("\n\n")
    chunks: list[str] = []
    current_chunk: list[str] = []
    current_len = 0

    for para in paragraphs:
        if current_len + len(para) > max_chars and current_chunk:
            chunks.append("\n\n".join(current_chunk))
            current_chunk = []
            current_len = 0
        current_chunk.append(para)
        current_len += len(para)

    if current_chunk:
        chunks.append("\n\n".join(current_chunk))

    return chunks


def extract_text(task: tuple[str, str]) -> tuple[str, list[str]]:
    """Extract (body, sections) for one (filepath, file_format) task."""
    filepath, file_format = task
    if file_format == "docx":
        return extract_docx_text(filepath)
    if file_format == "xlsx":
        return extract_xlsx_text(filepath)
    return "", []
